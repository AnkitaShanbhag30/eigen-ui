import os, json, time, zipfile
from typing import Dict
from .llm import generate_json
from .templates_loader import get_env
from .html_tokens import default_tokens
from .imgfm import gen_hero_image
from .util import ensure_dir, public_url
from docx import Document
from docx.shared import Pt
from markdown_it import MarkdownIt


def _font_links(heading, body):
	# Basic Google Fonts includes; safe even if font not present
	def link_for(name):
		q = name.replace(" ", "+")
		return f'<link href="https://fonts.googleapis.com/css2?family={q}:wght@400;600;700&display=swap" rel="stylesheet"/>'
	if heading == body:
		return link_for(heading)
	return link_for(heading) + link_for(body)


def outline_for(template: str, brand: Dict, x: str, y: str, z: str, w: str, cta: str) -> Dict:
	"""Generate content outline using a brand-agnostic multi-step flow.

	The flow replicates a planned, sequential prompting approach while remaining
	fully compatible with existing CLI and tests. It proceeds as follows:
	  1) Plan sections for the final asset (brand-agnostic)
	  2) For each section, generate content with brand JSON and explicit guidance
	  3) Assemble a final outline with headline/subhead/cta

	Falls back to the previous brand-aware system and finally to a static
	fallback if any step fails. Uses app.llm.generate_json so tests can mock it.
	"""
	try:
		# Attempt multi-step prompting orchestrated locally in this module
		outline = build_outline_multistep(template, brand, x, y, z, w, cta)
		if outline and isinstance(outline, dict) and outline.get("sections"):
			return outline
	except Exception as e:
		print(f"Warning: Multi-step prompting failed: {e}")

	# Fallback to the brand-aware outline generator (single-shot with UI insights)
	try:
		from .design import DesignAdvisorService
		design_advisor = DesignAdvisorService()

		if isinstance(brand, dict):
			from .brand import BrandIdentity
			brand_obj = BrandIdentity(**brand)
		else:
			brand_obj = brand

		outline = design_advisor.generate_content_outline(template, x, y, z, w, cta, brand_obj)
		if outline and isinstance(outline, dict):
			return outline
		else:
			print(f"Warning: Enhanced content generation returned invalid result: {outline}")
			return _get_fallback_outline(template, x, y, z, cta)
	except Exception as e:
		print(f"Warning: Enhanced content generation failed: {e}")
		return _get_fallback_outline(template, x, y, z, cta)


def _get_fallback_outline(template: str, x: str, y: str, z: str, cta: str) -> Dict:
	"""Fallback outline when LLM fails"""
	return {
		"headline": f"{x} - {y}",
		"subhead": f"Transform your {z} experience with our innovative solution",
		"sections": [
			{
				"title": "Why Choose Us",
				"bullets": [
					f"Built specifically for {z}",
					f"Delivers on {y}",
					"Proven results and customer satisfaction"
				]
			},
			{
				"title": "Key Benefits",
				"bullets": [
					"Streamlined workflow",
					"Increased efficiency",
					"Better outcomes"
				]
			}
		],
		"cta": cta or "Get Started Today"
	}


def _website_excerpt(brand: Dict) -> str:
	"""Fetch and condense the brand website content for LLM context."""
	try:
		website = (brand or {}).get("website")
		if not website:
			return ""
		from .scrape import fetch_html, visible_text_samples
		html = fetch_html(website, 5000) or ""
		if not html:
			return ""
		samples = visible_text_samples(html) or []
		text = " ".join(samples[:20])
		return (text[:2000] + "…") if len(text) > 2000 else text
	except Exception:
		return ""


def _synthesize_section_fallback(content_type: str, title_hint: str, brand: Dict, x: str, y: str, z: str, cta: str) -> Dict:
	"""Create a non-generic fallback for a section when LLM content is unavailable."""
	keywords = [kw for kw in (brand.get("keywords") or []) if isinstance(kw, str)]
	primary = (brand.get("colors", {}) or {}).get("primary")
	def pick_keywords(n: int) -> list:
		seen = set()
		picked = []
		for kw in keywords:
			k = kw.strip()
			if not k or k.lower() in {"menu", "search", "pages", "channel"}: continue
			if k.lower() in seen: continue
			seen.add(k.lower())
			picked.append(k)
			if len(picked) >= n: break
		return picked

	ctype = (content_type or "general").lower()
	if ctype == "hero":
		return {"title": title_hint or "Hero", "bullets": [], "content_type": "hero", "layout_style": "hero"}
	if ctype == "about":
		return {
			"title": title_hint or "About",
			"description": y or "What we deliver and why it matters",
			"bullets": [f"What: {x}", f"Why: {y}", f"For: {z}"],
			"content_type": "about",
			"layout_style": "text",
			"blocks": [
				{"type": "paragraph", "text": (y or "")[:240]},
				{"type": "bullets", "items": [f"What: {x}", f"Why: {y}", f"For: {z}"]},
			],
		}
	if ctype == "features":
		feats = pick_keywords(3) or ["Personalization", "Insights", "Optimization"]
		return {
			"title": title_hint or "Key Capabilities",
			"description": "Highlights that support the value proposition",
			"bullets": [f"{f} built-in" for f in feats],
			"content_type": "features",
			"layout_style": "grid",
			"blocks": [
				{"type": "statement", "text": "What sets us apart"},
				{"type": "columns", "columns": [f for f in feats]},
			],
		}
	if ctype == "testimonials":
		return {
			"title": title_hint or "What Our Clients Say",
			"description": "Real outcomes from real brands",
			"bullets": ["Proven conversion lifts", "Faster go-live", "Delighted customers"],
			"content_type": "testimonials",
			"layout_style": "list",
			"blocks": [
				{"type": "statement", "text": "Trusted by teams like yours"},
				{"type": "bullets", "items": ["Proven conversion lifts", "Faster go-live", "Delighted customers"]},
			],
		}
	if ctype == "contact":
		return {
			"title": title_hint or "Contact",
			"description": "We'd love to talk",
			"bullets": ["Book a demo", "Talk to our team", "Get a tailored walkthrough"],
			"content_type": "contact",
			"layout_style": "text",
			"blocks": [
				{"type": "paragraph", "text": "We'd love to show you a tailored walkthrough."},
				{"type": "bullets", "items": ["Book a demo", "Talk to our team", "Get a tailored walkthrough"]},
			],
		}
	if ctype == "cta":
		cta_text = cta or "Get Started"
		return {
			"title": title_hint or cta_text,
			"description": "",
			"bullets": [cta_text],
			"content_type": "cta",
			"layout_style": "text",
			"blocks": [
				{"type": "statement", "text": cta_text}
			],
		}
	# Generic fallback
	return {
		"title": title_hint or content_type.title(),
		"description": y or "",
		"bullets": [f"For {z}", f"Focus: {x}"],
		"content_type": content_type,
		"layout_style": "text",
		"blocks": [
			{"type": "paragraph", "text": (y or "").strip()[:200]}
		],
	}


def _safe_brand_snapshot(brand: Dict) -> Dict:
	"""Return a compact, safe-to-serialize snapshot of brand JSON for prompts."""
	try:
		# Keep only the fields the LLM needs to respect styling and tone
		colors = brand.get("colors", {}) if isinstance(brand, dict) else {}
		typography = brand.get("typography", {}) if isinstance(brand, dict) else {}
		return {
			"name": brand.get("name", ""),
			"website": brand.get("website", ""),
			"tagline": brand.get("tagline", ""),
			"tone": brand.get("tone", ""),
			"keywords": brand.get("keywords", [])[:12],
			"colors": {
				"primary": colors.get("primary"),
				"secondary": colors.get("secondary"),
				"accent": colors.get("accent"),
				"muted": colors.get("muted"),
			},
			"typography": {
				"heading": typography.get("heading", "Inter"),
				"body": typography.get("body", "Inter"),
			},
		}
	except Exception:
		return {}


def _default_section_plan_for(template: str) -> list[Dict]:
	"""Default section plan used when planning fails (brand-agnostic)."""
	# A simple, brand-agnostic one-pager structure that mirrors the example flow
	base = [
		{"content_type": "hero", "title_hint": "Hero"},
		{"content_type": "about", "title_hint": "About"},
		{"content_type": "features", "title_hint": "Features"},
		{"content_type": "testimonials", "title_hint": "Testimonials"},
		{"content_type": "contact", "title_hint": "Contact"},
		{"content_type": "cta", "title_hint": "Get Started"},
	]
	# Could vary per template in the future
	return base


def plan_sections_for_template(template: str, brand: Dict, x: str, y: str, z: str, w: str, website_excerpt: str = "") -> Dict:
	"""Plan proposed layout + ordered sections with themes via LLM.

	Expected JSON (preferred):
	{
	  "layout": {
		"variant": "feature-rail-3up|two-column|hero-left-cta|hero-right-cta|simple",
		"has_hero": true,
		"grid": {"columns": 1|2|3},
		"density": "airy|balanced|compact",
		"order": ["hero","about","features","testimonials","contact","cta"]
	  },
	  "sections": [
		{"content_type":"hero","title_hint":"...","theme":"...","layout_style":"hero","importance":"high"},
		{"content_type":"about","title_hint":"...","theme":"...","layout_style":"text"}
	  ]
	}
	If the model returns a complete outline instead, we accept it.
	"""
	print("🧭 Planning sections (multi-step)...")
	safe_brand = _safe_brand_snapshot(brand)
	system = (
		"PHASE: PLAN\n"
		"You propose the layout AND the ordered sections for a brand-agnostic asset, and you may also produce a complete outline.\n"
		"Use BOTH the brand JSON and the website excerpt for substance, tone, and specificity.\n"
		"Quality guardrails: ensure each section serves a distinct purpose (no repetition),\n"
		"uses specific, benefit-driven language, avoids placeholders, and yields a scannable hierarchy.\n"
		"Ensure DIVERSITY of content mix across sections (some short paragraphs, some bullet lists, some short statements, and at least one columns/grid presentation).\n"
		"Prefer accessible headings and contrast.\n"
		"Return either: (A) a complete outline JSON {headline, subhead, sections, cta} where each section includes a 'blocks' array,\n"
		"or (B) a planning JSON with keys 'layout' and 'sections' as described below.\n"
		"- layout: {variant, has_hero, grid:{columns}, density, order[]}\n"
		"- sections: array of {content_type, title_hint, theme, layout_style, importance, block_types?[]}\n"
		"Blocks schema (for outline case): each section.blocks[] items are one of: {type:'paragraph', text} | {type:'bullets', items[]} | {type:'statement', text} | {type:'columns', columns[]}"
	)
	user = json.dumps({
		"template": template,
		"brand": safe_brand,
		"requirements": {
			"what": x,
			"why": y,
			"audience": z,
			"context": w,
		},
		"website_excerpt": website_excerpt
	}, indent=2)

	try:
		response = generate_json(system, user)
		if isinstance(response, dict) and response:
			return response
	except Exception as e:
		print(f"Section planning failed: {e}")
	return {
		"layout": {"variant": "simple", "has_hero": False, "grid": {"columns": 1}, "density": "balanced", "order": [s["content_type"] for s in _default_section_plan_for(template)]},
		"sections": _default_section_plan_for(template)
	}


def generate_section_content(section_plan: Dict, brand: Dict, template: str, x: str, y: str, z: str, w: str, cta: str, website_excerpt: str = "") -> Dict:
	"""Generate content for a single section using brand JSON and requirements."""
	content_type = section_plan.get("content_type", "general")
	title_hint = section_plan.get("title_hint", "")
	theme = section_plan.get("theme", "")
	print(f"✍️  Generating section: {content_type}...")
	safe_brand = _safe_brand_snapshot(brand)

	system = (
		"PHASE: SECTION\n"
		"You are generating ONE section (content) for a brand-agnostic asset.\n"
		"Use BOTH the brand JSON and the website excerpt for substance, tone, accurate details, and specificity.\n"
		"Avoid generic titles like 'Features' or 'Benefits'—use concrete, compelling phrasing.\n"
		"Quality guardrails: content must be specific and non-repetitive; avoid placeholders; ensure scannability and clarity; if content_type is 'hero', avoid any text-over-image instructions.\n"
		"Return ONLY JSON with keys: {title, content_type, layout_style, description?, bullets?, blocks?}.\n"
		"Blocks provide a diverse content mix. Each block is one of: {type:'paragraph', text}, {type:'bullets', items[]}, {type:'statement', text}, {type:'columns', columns[]}\n"
		"Include 1-3 blocks per section and vary block types across sections."
	)
	user = json.dumps({
		"template": template,
		"brand": safe_brand,
		"requirements": {
			"what": x,
			"why": y,
			"audience": z,
			"context": w,
			"cta": cta,
		},
		"section_request": {
			"content_type": content_type,
			"title_hint": title_hint,
			"theme": theme,
		},
		"website_excerpt": website_excerpt
	}, indent=2)

	try:
		response = generate_json(system, user)
		if isinstance(response, dict) and response.get("title"):
			# Normalize
			out = {
				"title": response.get("title", title_hint or content_type.title()),
				"description": response.get("description") or "",
				"bullets": response.get("bullets") or [],
				"content_type": response.get("content_type", content_type),
				"layout_style": response.get("layout_style", "text"),
				"blocks": response.get("blocks") or []
			}
			return out
	except Exception as e:
		print(f"Section generation error: {e}")

	# Fallback minimal section
	return _synthesize_section_fallback(content_type, title_hint, brand, x, y, z, cta)


def _compose_headline_and_cta(brand: Dict, x: str, y: str, z: str, cta: str) -> Dict:
	"""Generate or infer headline/subhead/cta. Uses LLM but tolerates failure."""
	print("🧱 Composing headline and CTA...")
	safe_brand = _safe_brand_snapshot(brand)
	system = (
		"PHASE: SUMMARY\n"
		"Create a punchy headline (5-8 words), a concise subhead, and a CTA line.\n"
		"Quality guardrails: headline must be benefit-driven and specific (avoid generic claims);\n"
		"subhead should clarify value succinctly; CTA should be action + specific value (avoid 'Learn More' unless truly best).\n"
		"Respect tone and keywords from the brand JSON. Return ONLY JSON: {headline, subhead, cta}."
	)
	user = json.dumps({
		"brand": safe_brand,
		"what": x,
		"why": y,
		"audience": z,
		"cta_pref": cta,
	}, indent=2)
	try:
		response = generate_json(system, user)
		if isinstance(response, dict) and response.get("headline"):
			return {
				"headline": response.get("headline"),
				"subhead": response.get("subhead", ""),
				"cta": response.get("cta", cta or "Get Started"),
			}
	except Exception as e:
		print(f"Headline/CTA composition failed: {e}")
	# Fallbacks guided by inputs
	return {
		"headline": f"{x}" if x else (brand.get("tagline") or brand.get("name") or "Welcome"),
		"subhead": f"{y}" if y else (brand.get("description") or ""),
		"cta": cta or "Get Started",
	}


def build_outline_multistep(template: str, brand: Dict, x: str, y: str, z: str, w: str, cta: str) -> Dict:
	"""End-to-end multi-step outline builder orchestrating planning and sections."""
	# Gather website context once
	website_excerpt = _website_excerpt(brand)
	# Step 1: Plan
	plan = plan_sections_for_template(template, brand, x, y, z, w, website_excerpt)
	# If caller accidentally returned a full outline, accept it
	if plan.get("sections") and any(isinstance(s, dict) and ("title" in s or "bullets" in s) for s in plan["sections"]):
		# Looks like a full outline already
		if "headline" in plan or "cta" in plan:
			return plan
	sections_plan = plan.get("sections") or _default_section_plan_for(template)

	# If a layout order is provided, re-order sections_plan to match
	layout_info = plan.get("layout") or {}
	order = layout_info.get("order")
	if order and isinstance(order, list):
		order_index = {ct: i for i, ct in enumerate(order)}
		sections_plan = sorted(sections_plan, key=lambda s: order_index.get(s.get("content_type", "zzzz"), 999))

	# Step 2: Generate each section
	generated_sections = []
	for sp in sections_plan:
		try:
			generated_sections.append(generate_section_content(sp, brand, template, x, y, z, w, cta, website_excerpt))
		except Exception as e:
			print(f"Section generation fallback due to error: {e}")
			generated_sections.append({
				"title": sp.get("title_hint", sp.get("content_type", "Section")),
				"bullets": [],
				"content_type": sp.get("content_type", "general"),
				"layout_style": "text",
				"blocks": []
			})

	# Step 3: Headline/Subhead/CTA
	header = _compose_headline_and_cta(brand, x, y, z, cta)

	# Assemble outline
	outline = {
		"headline": header.get("headline", ""),
		"subhead": header.get("subhead", ""),
		"sections": generated_sections,
		"cta": header.get("cta", cta or "Get Started"),
		"meta": {
			"planned_sections": [sp.get("content_type") for sp in sections_plan],
			"proposed_layout": layout_info,
			"template": template,
		},
	}
	return outline


def polish_outline(o: Dict, brand: Dict) -> Dict:
	"""Polish the outline with better content"""
	if not o or not isinstance(o, dict):
		print(f"Warning: Invalid outline received: {o}")
		return o
	
	# For now, return the original outline to avoid LLM issues
	# In the future, this could be enhanced with LLM polishing
	return o

from typing import Optional


def render_html(template_key: str, brand: Dict, tokens: Dict, outline: Dict, hero_url: Optional[str]):
	"""Render HTML using the new template system with proper custom_data structure"""
	try:
		from .renderer import render_template_with_brand
		
		# Prepare custom data from outline and other parameters
		custom_data = {
			'title': outline.get("headline", ""),
			'subtitle': outline.get("subhead", ""),
			'cta': outline.get("cta", ""),
			'hero_url': hero_url,
			'outline': outline,
			'font_links': _font_links(brand["typography"]["heading"], brand["typography"]["body"])
		}
		
		# Use the new renderer system
		html = render_template_with_brand(template_key, brand, custom_data)
		return html
		
	except Exception as e:
		# Fallback to old system if new renderer fails
		print(f"⚠️  New renderer failed, falling back to old system: {e}")
		env = get_env()
		
		# Handle new channel structure
		template_path = template_key
		if template_key == 'story':
			template_path = 'story/story-highlights.html.j2'
		elif template_key == 'linkedin':
			template_path = 'linkedin/li-product-announcement.html.j2'
		else:
			template_path = f"{template_key}.html.j2"
		
		try:
			tpl = env.get_template(template_path)
		except Exception as template_error:
			print(f"⚠️  Template not found: {template_path}, trying fallback...")
			# Try the base template as last resort
			tpl = env.get_template("onepager.html.j2")
		
		# Ensure outline has minimum content
		if not outline.get("headline"):
			outline["headline"] = f"Welcome to {brand.get('name', 'Our Platform')}"
		if not outline.get("subhead"):
			outline["subhead"] = f"Discover how we can help you achieve your goals"
		if not outline.get("sections"):
			outline["sections"] = [
				{
					"title": "Get Started",
					"bullets": ["Simple setup", "Quick results", "Expert support"]
				}
			]
		if not outline.get("cta"):
			outline["cta"] = "Learn More"
		
		html = tpl.render(
			brand=brand,
			tokens=tokens,
			outline=outline,
			hero_url=hero_url,
			title=outline.get("headline",""),
			font_links=_font_links(brand["typography"]["heading"], brand["typography"]["body"])
		)
		return html

def write_pdf(html: str, out_path: str) -> bool:
    import os
    if os.getenv("PREFERRED_PDF_ENGINE","weasyprint") != "weasyprint":
        return False
    try:
        from weasyprint import HTML
        HTML(string=html).write_pdf(out_path)
        return True
    except Exception:
        return False

def write_docx(outline: Dict, brand: Dict, out_path: str):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = brand["typography"]["body"]
    style.font.size = Pt(11)
    h1 = doc.add_heading(outline.get("headline",""), 0)
    h1.style.font.name = brand["typography"]["heading"]
    if outline.get("subhead"):
        doc.add_paragraph(outline["subhead"])
    for s in outline.get("sections",[]):
        doc.add_heading(s.get("title",""), level=1)
        for b in s.get("bullets",[]):
            doc.add_paragraph(b, style=None).style = doc.styles['List Bullet']
    doc.add_paragraph(f"CTA: {outline.get('cta','')}")
    doc.save(out_path)

def zip_outputs(paths, out_zip):
    with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as z:
        for p in paths:
            if p and os.path.exists(p): z.write(p, arcname=os.path.basename(p))

def generate_assets(slug: str, brand: Dict, template: str, x:str,y:str,z:str,w:str,cta:str, hero_mode:str="skip", force_html:bool=False):
    ts = int(time.time())
    drafts_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "data", "drafts", slug))
    assets_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "data", "assets", slug))
    ensure_dir(drafts_dir); ensure_dir(assets_dir)

    # Generate content outline
    print(f"🤖 Generating content outline for {template}...")
    outline = outline_for(template, brand, x, y, z, w, cta)
    
    # Ensure CTA is set
    if cta and cta.strip():
        outline["cta"] = cta.strip()
    
    # Polish the outline
    print(f"✨ Polishing content...")
    outline = polish_outline(outline, brand)
    
    # Debug output
    print(f"📝 Generated outline: {outline.get('headline', 'N/A')} | {outline.get('cta', 'N/A')} | {len(outline.get('sections', []))} sections")
    
    # build tokens
    heading = brand.get("typography",{}).get("heading","Inter")
    body = brand.get("typography",{}).get("body","Inter")
    tokens = default_tokens(brand.get("colors",{}), heading, body)

    hero_url = None
    if hero_mode.startswith("path:"):
        local = hero_mode.split("path:",1)[1]
        hero_url = public_url(local)
    elif hero_mode == "auto" or os.getenv("IMAGE_PROVIDER","" ).lower() == "fal" or os.getenv("FAL_API_KEY"):
        tone = brand.get('tone','') or 'professional'
        prompt = brand.get("heroBrief","") or f"abstract clean hero image, {tone}"
        # Prefer remote URL via Fal when available
        try:
            from .imgfm import generate_hero_image_remote_url
            remote = generate_hero_image_remote_url(prompt, brand.get('keywords', []), [brand.get('colors',{}).get('primary')])
        except Exception:
            remote = None
        if remote:
            hero_url = remote
        else:
            hero_path = os.path.join(assets_dir, f"hero_{ts}.png")
            if gen_hero_image(prompt, hero_path):
                hero_url = public_url(hero_path)

    # If template expects a hero and we still don't have one, error out instead of falling back
    if template == 'onepager' and not hero_url:
        raise RuntimeError("Hero image is required for onepager but was not provided or generated. Set --hero auto or --hero path:./image.png, or configure IMAGE_PROVIDER and FAL_API_KEY.")

    html = render_html(template, brand, tokens, outline, hero_url)
    html_path = os.path.join(drafts_dir, f"{slug}-{template}-{ts}.html")
    
    # Check if Dyad templates are available and use SSR renderer
    html_for_pdf = html  # Default to original HTML
    try:
        import sys
        from pathlib import Path
        sys.path.insert(0, str(Path(__file__).parent.parent))
        from renderer.resolve_templates import resolve_templates
        import subprocess
        
        sel = resolve_templates()
        
        if sel["engine"] == "react" and not force_html:
            print("🎨 Using Dyad React SSR renderer...")
            entry = sel["entry"]
            # Ensure brand json exists at data/brands/{slug}.json
            brand_json_path = f"data/brands/{slug}.json"
            if not os.path.exists(brand_json_path):
                # Create brand json if it doesn't exist
                with open(brand_json_path, 'w') as f:
                    json.dump(brand, f, indent=2, default=str)
                print(f"💾 Created brand config: {brand_json_path}")
            
            try:
                # Pass user requirements to SSR renderer
                cmd = [
                    "pnpm", "ssr",
                    "--brand", slug,
                    "--entry", entry,
                    "--out", html_path,
                    "--props", brand_json_path
                ]
                
                # Add user requirements if available
                if x: cmd.extend(["--x", x])
                if y: cmd.extend(["--y", y])
                if z: cmd.extend(["--z", z])
                if w: cmd.extend(["--w", w])
                if cta: cmd.extend(["--cta", cta])
                
                # Quietly capture output to avoid noisy Node stack traces on failure
                # Inject generated hero URL and outline into props for SSR consumption
                try:
                    if os.path.exists(brand_json_path):
                        with open(brand_json_path, 'r') as f:
                            brand_payload = json.load(f)
                    else:
                        brand_payload = dict(brand)
                    if hero_url:
                        gi = brand_payload.get('generatedImages') or {}
                        gi['hero'] = hero_url
                        brand_payload['generatedImages'] = gi
                    # Provide the generated outline directly to SSR
                    brand_payload['generatedOutline'] = outline
                    with open(brand_json_path, 'w') as f:
                        json.dump(brand_payload, f, indent=2, default=str)
                except Exception:
                    pass

                subprocess.run(
                    cmd,
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True
                )
                print(f"✅ SSR rendered to: {html_path}")
                
                # Read the generated HTML for PDF conversion
                with open(html_path, "r", encoding="utf-8") as f:
                    html_for_pdf = f.read()
            except subprocess.CalledProcessError as e:
                print(f"⚠️ SSR renderer failed, falling back to HTML: {e}")
                with open(html_path, "w", encoding="utf-8") as f: 
                    f.write(html)
                html_for_pdf = html
        else:
            # If SSR engine is React but dependencies are missing, skip SSR to avoid errors
            if sel.get("engine") == "react" and not force_html:
                node_esbuild = os.path.join(os.getcwd(), "node_modules", "esbuild")
                if not os.path.exists(node_esbuild):
                    print("ℹ️  SSR dependencies not installed (esbuild missing). Skipping SSR. Run 'pnpm install' to enable.")
            # Use existing HTML writer
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html)
    except Exception as e:
        print(f"⚠️ SSR integration failed, falling back to HTML: {e}")
        with open(html_path, "w", encoding="utf-8") as f: 
            f.write(html)

    pdf_path = os.path.join(drafts_dir, f"{slug}-{template}-{ts}.pdf")
    pdf_ok = write_pdf(html_for_pdf, pdf_path)
    if not pdf_ok: pdf_path = None

    docx_path = os.path.join(drafts_dir, f"{slug}-{template}-{ts}.docx")
    write_docx(outline, brand, docx_path)

    zip_path = os.path.join(drafts_dir, f"{slug}-{template}-{ts}.zip")
    zip_outputs([p for p in [html_path, pdf_path, docx_path] if p], zip_path)

    return {
        "outline": outline,
        "paths": {
            "html": html_path,
            "pdf": pdf_path,
            "docx": docx_path,
            "zip": zip_path
        },
        "public": {
            "html": public_url(html_path),
            "pdf": public_url(pdf_path) if pdf_path else None,
            "docx": public_url(docx_path),
            "zip": public_url(zip_path),
            "hero": hero_url
        },
        "tokens": tokens
    } 